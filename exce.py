import openpyxl
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open Excel workbook and select worksheet
wb = openpyxl.load_workbook('example.xlsx')
ws = wb.active

# Define formatting styles
title_style = 'Title'
header_style = 'Header'
body_style = 'Body'

# Open blank Word document
template = Document('template.docx')

# Add title
title = template.add_heading('Title', level=1)
title.style = title_style

# Add table
table = template.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Add headers
header_cells = table.rows[0].cells
header_cells[0].text = 'Header 1'
header_cells[1].text = 'Header 2'
header_cells[2].text = 'Header 3'
for cell in header_cells:
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].style = header_style

# Add data
for row in ws.iter_rows():
    row_cells = [cell.value for cell in row]
    row_cells = [str(cell) if cell is not None else '' for cell in row_cells]
    new_row = table.add_row().cells
    for i, cell_value in enumerate(row_cells):
        new_row[i].text = cell_value
        new_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        new_row[i].paragraphs[0].style = body_style

# Save modified Word document
template.save('output.docx')
